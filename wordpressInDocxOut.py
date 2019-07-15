import xml.etree.ElementTree as ET
from html.parser import HTMLParser
from docx import Document
import datetime
import string
import requests
import pathlib
import shutil
from docx.shared import Inches

class MyHTMLParser(HTMLParser):
 
    def __init__(self, _lastTitle, _lastDate):
        self.doc = Document()
        self.doc.add_heading(_lastTitle, 0)
        self.doc.add_heading(_lastDate, level=1)
        self.imageCounter = 0
        self.imageStack = []
        self.treeStack = []
        super().__init__()

    def handle_starttag(self, tag, attrs):
        print(f"Encountered a start tag:{tag} , {attrs}")
        if (tag == 'img'):
            for attr in attrs:
                if attr[0] == 'src':
                    print(attr[1])
                    self.imageCounter = self.imageCounter + 1
                    extention = pathlib.Path(attr[1]).suffix
                    fileName = './exports/' + str(self.imageCounter) + extention
                    print(fileName)
                    url = attr[1]
                    response = requests.get(url, stream=True)
                    with open(fileName, 'wb') as out_file:
                        shutil.copyfileobj(response.raw, out_file)
                    del response
                    #with open(fileName, 'wb') as f:
                    #    f.write(requests.get(attr[1]).content)
                    self.imageStack.append(fileName)
        self.treeStack.append(tag)
        pass

    def handle_endtag(self, tag):
        print("Encountered an end tag :", tag)
        tagtype = self.treeStack.pop()
        pass

    def handle_data(self, data):
        strippedString = data.translate(str.maketrans('', '', string.whitespace))
        print(f"Encountered some data  :{len(data)}{len(strippedString)}{data}")
        if (len(strippedString) > 0):
            if (len(self.treeStack)>0):
                tagtype = self.treeStack[len(self.treeStack)-1]
                if (tagtype == 'p'):
                    if (len(self.treeStack) > 1):
                        innertagtype = self.treeStack[len(self.treeStack)-2]
                        if innertagtype == 'blockquote':
                            textToWrite = '\"' + data + '\"'
                            self.doc.add_paragraph(textToWrite)
                    else:
                        self.doc.add_paragraph(data)
                if (tagtype == 'h1'):
                    self.doc.add_heading(data, level=1)
                if (tagtype == 'h2'):
                    self.doc.add_heading(data, level=2)
                if (tagtype == 'h3'):
                    self.doc.add_heading(data, level=3)
                if (tagtype == 'div'):
                    #self.doc.add_paragraph(data)
                    innertagtype = self.treeStack[len(self.treeStack)-2]
                    if (innertagtype == 'figure'):
                        self.doc.add_paragraph(data)
                if (tagtype == 'a'):
                    #self.doc.add_paragraph(data)
                    innertagtype = self.treeStack[len(self.treeStack)-2]
                    if (innertagtype == 'p'):
                        self.doc.add_paragraph(data)
                if (tagtype == 'li'):
                    #self.doc.add_paragraph(data)
                    innertagtype = self.treeStack[len(self.treeStack)-2]
                    if (innertagtype == 'ul'):
                        self.doc.add_paragraph(data, style='List Bullet')
                if (tagtype == 'img'):
                    fileName = self.imageStack[len(self.treeStack)-1]
                    print(f"write {fileName} to the doc")
                    self.doc.add_picture(fileName, width=Inches(1.25))
            else:
                self.doc.add_paragraph(data)
        else:
            if (len(self.imageStack)>0):
                fileName = self.imageStack.pop()
                print(f"write {fileName} to the doc")
                self.doc.add_picture(fileName, width=Inches(1.25))
   
    def write_document(self, title):
        self.doc.save('./exports/'+title+'.docx')

tree = ET.parse('researchandideasdiary.WordPress.2019-07-15.xml')
root = tree.getroot()
print(root.tag)
print(root.items())
channel = root.getchildren()[0]
listOfTags = []
for child in channel.getchildren():
    print(child.tag, child.attrib)
    if child.tag == 'item':
        lastTitle = ''
        lastDate = ''
        lastGuid = ''
        for postdata in child.getchildren():
            listOfTags.append(postdata.tag)
            if (postdata.tag == 'title'):
                print(postdata.text)
                lastTitle = postdata.text
            if (postdata.tag == 'pubDate'):
                print(postdata.text)
                lastDate = postdata.text
            if (postdata.tag == 'guid'):
                print(postdata.text)
                lastGuid = postdata.text
            if postdata.tag == '{http://purl.org/rss/1.0/modules/content/}encoded':
                datestring = 'draft'
                if (lastDate != 'Mon, 30 Nov -0001 00:00:00 +0000'):
                    date_time_obj = datetime.datetime.strptime(lastDate, '%a, %d %b %Y %H:%M:%S %z')
                    datestring = str(date_time_obj.date())
                print(datestring)
                parser = MyHTMLParser(lastTitle, datestring)
                parser.feed(str(postdata.text))
                parser.write_document('Exp'+datestring+lastTitle)

listOfUniqueTags = list(set(listOfTags))
for tag in listOfUniqueTags:
    print(tag)

                